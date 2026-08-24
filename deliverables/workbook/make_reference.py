"""Build the pandoc reference document carrying IU formatting rules.

Pandoc's own default reference.docx is used as the base because it is the only
document guaranteed to define every style id the docx writer emits
(BodyText, FirstParagraph, Compact, ImageCaption, Verbatim, ...). The IU
template ships a different, smaller style set, so basing the reference on it
would silently drop pandoc output into unstyled defaults.

The formatting values below come from "Guidelines for structuring and
formatting academic papers" (IU), which is the graded specification.
"""

from __future__ import annotations

import copy
import subprocess
import sys
from pathlib import Path

import docx
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.shared import Cm

HERE = Path(__file__).resolve().parent
OUT = HERE / "assets" / "reference.docx"

TWIPS_PER_PT = 20
A4_WIDTH_TWIPS = 11906
A4_HEIGHT_TWIPS = 16838

BODY_FONT = "Arial"
MONO_FONT = "Consolas"
BLACK = "000000"

#: styleId -> (half-points, bold, alignment, space-before-pt, space-after-pt, line-spacing)
#: line spacing is expressed in 240ths of a line: 360 == 1.5 lines, 240 == single.
STYLE_SPEC: dict[str, dict] = {
    # Body text: Arial 11 pt, 1.5 spacing, justified, 6 pt between paragraphs.
    "Normal": dict(sz=22, jc="both", after=6, line=360),
    "BodyText": dict(sz=22, jc="both", after=6, line=360),
    "FirstParagraph": dict(sz=22, jc="both", after=6, line=360),
    "Compact": dict(sz=22, jc="both", after=6, line=360),
    "Author": dict(sz=22, jc="center", after=6, line=360),
    "Date": dict(sz=22, jc="center", after=6, line=360),
    "BlockText": dict(sz=22, jc="both", after=6, line=360),
    "Definition": dict(sz=22, jc="both", after=6, line=360),
    "DefinitionTerm": dict(sz=22, bold=True, jc="both", after=6, line=360),
    "Bibliography": dict(sz=22, jc="both", after=6, line=360, hanging=1.27),
    # Headings: left aligned, bold, black. Max three levels are permitted.
    "Heading1": dict(sz=32, bold=True, jc="left", before=12, after=12, line=360),
    "Heading2": dict(sz=28, bold=True, jc="left", before=12, after=6, line=360),
    "Heading3": dict(sz=22, bold=True, jc="left", before=12, after=6, line=360),
    "Heading4": dict(sz=22, bold=True, italic=True, jc="left", before=12, after=6, line=360),
    "TOCHeading": dict(sz=32, bold=True, jc="left", before=12, after=12, line=360),
    "AbstractTitle": dict(sz=28, bold=True, jc="left", before=12, after=6, line=360),
    "Abstract": dict(sz=22, jc="both", after=6, line=360),
    "Title": dict(sz=40, bold=True, jc="center", after=12, line=360),
    "Subtitle": dict(sz=28, jc="center", after=12, line=360),
    # Captions and footnotes: 10 pt, single spaced.
    "Caption": dict(sz=20, jc="both", after=6, line=240),
    "ImageCaption": dict(sz=20, jc="both", after=6, line=240),
    "TableCaption": dict(sz=20, jc="both", after=6, line=240),
    "FootnoteText": dict(sz=20, jc="both", after=0, line=240),
    "FootnoteBlockText": dict(sz=20, jc="both", after=0, line=240),
    # Figures are centred; their runs carry the image only.
    "Figure": dict(sz=22, jc="center", before=6, after=6, line=240),
    "CaptionedFigure": dict(sz=22, jc="center", before=6, after=6, line=240),
}

#: Character styles that must be monospaced so code stays legible and copyable.
MONO_STYLES = ("VerbatimChar",)


def _sub(parent, tag: str):
    """Return the child ``tag`` of ``parent``, creating it if absent.

    Word requires ``rPr``/``pPr`` children in schema order, so newly created
    property containers are inserted at the front of the element.
    """
    found = parent.find(qn(tag))
    if found is None:
        found = parent.makeelement(qn(tag), {})
        parent.insert(0, found)
    return found


def _set(parent, tag: str, **attrs) -> None:
    """Create or replace a simple property element such as ``<w:sz w:val="22"/>``."""
    existing = parent.find(qn(tag))
    if existing is not None:
        parent.remove(existing)
    el = parent.makeelement(qn(tag), {})
    for key, value in attrs.items():
        el.set(qn(f"w:{key}"), str(value))
    parent.append(el)


def _apply_font(rpr, *, name: str, sz: int | None = None,
                bold: bool | None = None, italic: bool | None = None) -> None:
    _set(rpr, "w:rFonts", ascii=name, hAnsi=name, cs=name, eastAsia=name)
    _set(rpr, "w:color", val=BLACK)
    if sz is not None:
        _set(rpr, "w:sz", val=sz)
        _set(rpr, "w:szCs", val=sz)
    if bold is not None:
        _set(rpr, "w:b", val="1" if bold else "0")
        _set(rpr, "w:bCs", val="1" if bold else "0")
    if italic is not None:
        _set(rpr, "w:i", val="1" if italic else "0")
        _set(rpr, "w:iCs", val="1" if italic else "0")


def _apply_paragraph(ppr, spec: dict) -> None:
    spacing = {"line": spec.get("line", 360), "lineRule": "auto",
               "before": int(spec.get("before", 0) * TWIPS_PER_PT),
               "after": int(spec.get("after", 0) * TWIPS_PER_PT)}
    _set(ppr, "w:spacing", **spacing)
    _set(ppr, "w:jc", val=spec.get("jc", "both"))
    # No first-line indent anywhere; the bibliography uses a hanging indent.
    if "hanging" in spec:
        _set(ppr, "w:ind", left=int(Cm(spec["hanging"]).twips),
             hanging=int(Cm(spec["hanging"]).twips))
    else:
        _set(ppr, "w:ind", firstLine=0, left=0)
    # Keep headings attached to the text they introduce.
    if spec.get("bold") and spec.get("jc") == "left":
        _set(ppr, "w:keepNext", val="1")


def patch_styles(document) -> None:
    """Rewrite the default document font and every style listed in STYLE_SPEC."""
    root = document.styles.element

    doc_defaults = root.find(qn("w:docDefaults"))
    if doc_defaults is not None:
        rpr_default = _sub(doc_defaults, "w:rPrDefault")
        _apply_font(_sub(rpr_default, "w:rPr"), name=BODY_FONT, sz=22)
        ppr_default = _sub(doc_defaults, "w:pPrDefault")
        _apply_paragraph(_sub(ppr_default, "w:pPr"), STYLE_SPEC["Normal"])

    by_id = {s.get(qn("w:styleId")): s for s in root.findall(qn("w:style"))}

    for style_id, spec in STYLE_SPEC.items():
        style = by_id.get(style_id)
        if style is None:
            continue
        _apply_font(_sub(style, "w:rPr"), name=BODY_FONT, sz=spec.get("sz"),
                    bold=spec.get("bold"), italic=spec.get("italic"))
        _apply_paragraph(_sub(style, "w:pPr"), spec)
        # Linked character styles must not contradict their paragraph style.
        linked = by_id.get(f"{style_id}Char")
        if linked is not None:
            _apply_font(_sub(linked, "w:rPr"), name=BODY_FONT, sz=spec.get("sz"),
                        bold=spec.get("bold"), italic=spec.get("italic"))

    for style_id in MONO_STYLES:
        style = by_id.get(style_id)
        if style is not None:
            _apply_font(_sub(style, "w:rPr"), name=MONO_FONT, sz=18)

    _add_source_code_style(root, by_id)


def _add_source_code_style(root, by_id: dict) -> None:
    """Add the paragraph style pandoc applies to fenced code blocks.

    Pandoc's stock reference document only defines the ``VerbatimChar`` run
    style, so without this the appendix listings inherit justified 1.5-spaced
    body text and become unreadable.
    """
    if "SourceCode" in by_id:
        return
    template = by_id.get("BodyText")
    if template is None:
        return
    style = copy.deepcopy(template)
    style.set(qn("w:styleId"), "SourceCode")
    for tag in ("w:name", "w:basedOn", "w:next", "w:link", "w:qFormat", "w:uiPriority"):
        found = style.find(qn(tag))
        if found is not None:
            style.remove(found)
    _set(style, "w:name", val="Source Code")
    _apply_font(_sub(style, "w:rPr"), name=MONO_FONT, sz=18)
    _apply_paragraph(_sub(style, "w:pPr"),
                     dict(jc="left", before=0, after=0, line=240))
    root.append(style)


def patch_section(document) -> None:
    """Set A4 paper with the mandated 2.00 cm margins on all four sides."""
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    sect_pr = section._sectPr
    _set(sect_pr, "w:pgSz", w=A4_WIDTH_TWIPS, h=A4_HEIGHT_TWIPS)
    _set(sect_pr, "w:pgMar",
         top=int(Cm(2).twips), right=int(Cm(2).twips),
         bottom=int(Cm(2).twips), left=int(Cm(2).twips),
         header=int(Cm(1.25).twips), footer=int(Cm(1.25).twips), gutter=0)


def add_page_number_footer(document) -> None:
    """Put a centred Arabic page number at the bottom of every page."""
    footer = document.sections[0].footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.text = ""
    ppr = _sub(paragraph._p, "w:pPr")
    _set(ppr, "w:jc", val="center")
    _set(ppr, "w:spacing", before=0, after=0, line=240, lineRule="auto")

    def run(*children):
        r = paragraph._p.makeelement(qn("w:r"), {})
        rpr = r.makeelement(qn("w:rPr"), {})
        _apply_font(rpr, name=BODY_FONT, sz=20)
        r.append(rpr)
        for child in children:
            r.append(child)
        paragraph._p.append(r)

    def el(tag, text=None, **attrs):
        e = paragraph._p.makeelement(qn(tag), {})
        for key, value in attrs.items():
            e.set(qn(f"w:{key}"), str(value))
        if text is not None:
            e.text = text
            e.set(qn("xml:space"), "preserve")
        return e

    run(el("w:fldChar", fldCharType="begin"))
    run(el("w:instrText", " PAGE \\* ARABIC "))
    run(el("w:fldChar", fldCharType="separate"))
    run(el("w:t", "1"))
    run(el("w:fldChar", fldCharType="end"))


def enable_hyphenation(document) -> None:
    """The formatting guidelines require hyphenation to be switched on."""
    settings = document.settings.element
    _set(settings, "w:autoHyphenation", val="true")
    _set(settings, "w:hyphenationZone", val=357)


def main() -> int:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    base = OUT.with_suffix(".base.docx")
    with base.open("wb") as fh:
        subprocess.run(["pandoc", "--print-default-data-file", "reference.docx"],
                       stdout=fh, check=True)

    document = docx.Document(str(base))
    patch_styles(document)
    patch_section(document)
    add_page_number_footer(document)
    enable_hyphenation(document)
    document.save(str(OUT))
    base.unlink()

    print(f"wrote {OUT.relative_to(HERE.parent.parent)}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
