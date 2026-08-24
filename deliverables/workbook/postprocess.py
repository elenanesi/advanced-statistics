"""Apply the two-part page numbering the IU formatting guidelines require.

Front matter is numbered in upper-case Roman numerals with no number shown on
the title page; the main text and appendices restart at Arabic 1. Word expresses
this with section breaks, which pandoc cannot emit, so the break is marked in the
Markdown with a sentinel paragraph and converted here.
"""

from __future__ import annotations

import copy
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn

MARKER = "%%SECTION-BREAK-ARABIC%%"

#: Headings that should keep their heading appearance but not be listed by the
#: table of contents field. A table of contents that lists itself is a defect.
TOC_EXCLUDED = {"table of contents"}

#: ``w:outlineLvl`` sits near the end of ``w:pPr``; these may legally follow it.
_PPR_TAIL = ("w:rPr", "w:sectPr", "w:pPrChange")


def _set(parent, tag: str, **attrs):
    existing = parent.find(qn(tag))
    if existing is not None:
        parent.remove(existing)
    el = parent.makeelement(qn(tag), {})
    for key, value in attrs.items():
        el.set(qn(f"w:{key}"), str(value))
    parent.append(el)
    return el


def _get_or_add_ppr(paragraph):
    ppr = paragraph._p.find(qn("w:pPr"))
    if ppr is None:
        ppr = paragraph._p.makeelement(qn("w:pPr"), {})
        paragraph._p.insert(0, ppr)
    return ppr


def split_sections(document) -> bool:
    """Turn the sentinel paragraph into a section break.

    The front-matter section inherits a copy of the document's final section
    properties, so page size, margins and the footer reference stay identical;
    only the numbering format differs.
    """
    body = document.element.body
    final_sect_pr = body.find(qn("w:sectPr"))
    if final_sect_pr is None:
        raise ValueError("document has no body-level section properties")

    marker = next((p for p in document.paragraphs if MARKER in p.text), None)
    if marker is None:
        return False

    front_sect_pr = copy.deepcopy(final_sect_pr)
    _set(front_sect_pr, "w:pgNumType", fmt="upperRoman", start="1")
    # Suppress the page number on the title page by giving the first page of the
    # front-matter section its own, empty footer.
    _set(front_sect_pr, "w:titlePg", val="1")

    ppr = _get_or_add_ppr(marker)
    ppr.append(front_sect_pr)
    for run in list(marker._p.findall(qn("w:r"))):
        marker._p.remove(run)

    _set(final_sect_pr, "w:pgNumType", fmt="decimal", start="1")
    return True


def add_blank_first_page_footer(document) -> None:
    """Give the front-matter section an empty first-page footer.

    Word looks for a ``first`` footer when ``titlePg`` is set; if none exists it
    falls back to the default footer and the title page would show "I".
    """
    section = document.sections[0]
    footer = section.first_page_footer
    footer.is_linked_to_previous = False
    for paragraph in footer.paragraphs:
        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)


def exclude_from_toc(document) -> int:
    """Drop selected headings out of the table of contents field.

    Outline level 9 is Word's "body text" level: the paragraph keeps the visual
    formatting of its heading style but the ``TOC \\o "1-3"`` field skips it.
    """
    count = 0
    for paragraph in document.paragraphs:
        if paragraph.text.strip().lower() not in TOC_EXCLUDED:
            continue
        if not paragraph.style.name.startswith("Heading"):
            continue
        ppr = _get_or_add_ppr(paragraph)
        existing = ppr.find(qn("w:outlineLvl"))
        if existing is not None:
            ppr.remove(existing)
        el = ppr.makeelement(qn("w:outlineLvl"), {})
        el.set(qn("w:val"), "9")
        tail = [child for child in ppr if child.tag in {qn(t) for t in _PPR_TAIL}]
        if tail:
            tail[0].addprevious(el)
        else:
            ppr.append(el)
        count += 1
    return count


def main() -> int:
    if len(sys.argv) != 2:
        print("usage: postprocess.py <workbook.docx>", file=sys.stderr)
        return 2
    path = Path(sys.argv[1])
    document = docx.Document(str(path))

    if not split_sections(document):
        print(f"warning: sentinel {MARKER} not found; page numbering unchanged",
              file=sys.stderr)
        return 1

    add_blank_first_page_footer(document)
    excluded = exclude_from_toc(document)
    document.save(str(path))

    sections = len(document.sections)
    print(f"page numbering applied: {sections} sections "
          "(front matter in Roman, main text restarting at Arabic 1); "
          f"{excluded} heading(s) excluded from the table of contents")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
